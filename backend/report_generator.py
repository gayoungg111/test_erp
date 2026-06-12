from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def _register_korean_font() -> str:
    font_paths = [
        "C:/Windows/Fonts/malgun.ttf",
        "C:/Windows/Fonts/malgunbd.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/System/Library/Fonts/AppleSDGothicNeo.ttc",
    ]
    for path in font_paths:
        try:
            pdfmetrics.registerFont(TTFont("Korean", path))
            return "Korean"
        except Exception:
            continue
    return "Helvetica"


def generate_word_report(summary: dict, data: list[dict]) -> bytes:
    doc = Document()

    title = doc.add_heading("ERP 분석 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}")
    doc.add_paragraph("")

    doc.add_heading("1. 요약", level=1)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Light Grid Accent 1"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "항목"
    hdr[1].text = "값"

    summary_items = [
        ("총 거래 건수", f"{summary.get('totalRecords', 0):,}건"),
        ("총 금액", f"{summary.get('totalAmount', 0):,.0f}원"),
        ("총 수량", f"{summary.get('totalQuantity', 0):,.0f}"),
        ("평균 거래 금액", f"{summary.get('avgAmount', 0):,.0f}원"),
        ("부서 수", f"{summary.get('departmentCount', 0)}개"),
        ("항목 수", f"{summary.get('itemCount', 0)}개"),
    ]
    date_range = summary.get("dateRange", {})
    if date_range:
        summary_items.append(
            ("분석 기간", f"{date_range.get('start', '-')} ~ {date_range.get('end', '-')}")
        )

    for label, value in summary_items:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph("")

    doc.add_heading("2. 부서별 분석", level=1)
    dept_data = summary.get("byDepartment", [])
    if dept_data:
        dept_table = doc.add_table(rows=1, cols=4)
        dept_table.style = "Light Grid Accent 1"
        headers = ["부서", "총금액", "총수량", "건수"]
        for i, h in enumerate(headers):
            dept_table.rows[0].cells[i].text = h
        for item in dept_data:
            row = dept_table.add_row().cells
            row[0].text = str(item.get("부서", ""))
            row[1].text = f"{item.get('총금액', 0):,.0f}"
            row[2].text = f"{item.get('총수량', 0):,.0f}"
            row[3].text = str(item.get("건수", 0))
    else:
        doc.add_paragraph("부서별 데이터가 없습니다.")

    doc.add_paragraph("")

    doc.add_heading("3. 항목별 Top 10", level=1)
    item_data = summary.get("byItem", [])
    if item_data:
        item_table = doc.add_table(rows=1, cols=3)
        item_table.style = "Light Grid Accent 1"
        headers = ["항목", "총금액", "총수량"]
        for i, h in enumerate(headers):
            item_table.rows[0].cells[i].text = h
        for item in item_data:
            row = item_table.add_row().cells
            row[0].text = str(item.get("항목", ""))
            row[1].text = f"{item.get('총금액', 0):,.0f}"
            row[2].text = f"{item.get('총수량', 0):,.0f}"
    else:
        doc.add_paragraph("항목별 데이터가 없습니다.")

    doc.add_paragraph("")

    doc.add_heading("4. 월별 추이", level=1)
    month_data = summary.get("byMonth", [])
    if month_data:
        month_table = doc.add_table(rows=1, cols=3)
        month_table.style = "Light Grid Accent 1"
        headers = ["월", "총금액", "건수"]
        for i, h in enumerate(headers):
            month_table.rows[0].cells[i].text = h
        for item in month_data:
            row = month_table.add_row().cells
            row[0].text = str(item.get("월", ""))
            row[1].text = f"{item.get('총금액', 0):,.0f}"
            row[2].text = str(item.get("건수", 0))

    doc.add_paragraph("")
    doc.add_heading("5. 분석 의견", level=1)
    opinion = doc.add_paragraph()
    opinion.add_run(_generate_insights(summary)).font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_pdf_report(summary: dict, data: list[dict]) -> bytes:
    buffer = io.BytesIO()
    font_name = _register_korean_font()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=18,
        alignment=1,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=13,
        spaceBefore=12,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
    )

    elements = []
    elements.append(Paragraph("ERP 분석 보고서", title_style))
    elements.append(
        Paragraph(
            f"생성일: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}",
            body_style,
        )
    )
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("1. 요약", heading_style))
    summary_rows = [
        ["항목", "값"],
        ["총 거래 건수", f"{summary.get('totalRecords', 0):,}건"],
        ["총 금액", f"{summary.get('totalAmount', 0):,.0f}원"],
        ["총 수량", f"{summary.get('totalQuantity', 0):,.0f}"],
        ["평균 거래 금액", f"{summary.get('avgAmount', 0):,.0f}원"],
        ["부서 수", f"{summary.get('departmentCount', 0)}개"],
        ["항목 수", f"{summary.get('itemCount', 0)}개"],
    ]
    date_range = summary.get("dateRange", {})
    if date_range:
        summary_rows.append(
            ["분석 기간", f"{date_range.get('start', '-')} ~ {date_range.get('end', '-')}"]
        )

    summary_table = Table(summary_rows, colWidths=[80 * mm, 80 * mm])
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
            ]
        )
    )
    elements.append(summary_table)
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("2. 부서별 분석", heading_style))
    dept_data = summary.get("byDepartment", [])
    if dept_data:
        dept_rows = [["부서", "총금액", "총수량", "건수"]]
        for item in dept_data:
            dept_rows.append(
                [
                    str(item.get("부서", "")),
                    f"{item.get('총금액', 0):,.0f}",
                    f"{item.get('총수량', 0):,.0f}",
                    str(item.get("건수", 0)),
                ]
            )
        dept_table = Table(dept_rows, colWidths=[40 * mm, 40 * mm, 40 * mm, 30 * mm])
        dept_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
                ]
            )
        )
        elements.append(dept_table)

    elements.append(Spacer(1, 12))
    elements.append(Paragraph("3. 분석 의견", heading_style))
    elements.append(Paragraph(_generate_insights(summary), body_style))

    doc.build(elements)
    return buffer.getvalue()


def _generate_insights(summary: dict) -> str:
    total = summary.get("totalAmount", 0)
    records = summary.get("totalRecords", 0)
    dept_data = summary.get("byDepartment", [])

    lines = [
        f"총 {records:,}건의 ERP 거래 데이터를 분석한 결과, 총 금액은 {total:,.0f}원입니다.",
    ]

    if dept_data:
        top_dept = dept_data[0]
        lines.append(
            f"부서별로는 '{top_dept.get('부서')}' 부서가 "
            f"{top_dept.get('총금액', 0):,.0f}원으로 가장 높은 비중을 차지합니다."
        )

    month_data = summary.get("byMonth", [])
    if len(month_data) >= 2:
        last = month_data[-1].get("총금액", 0)
        prev = month_data[-2].get("총금액", 0)
        if prev > 0:
            change = ((last - prev) / prev) * 100
            direction = "증가" if change >= 0 else "감소"
            lines.append(
                f"최근 월별 추이를 보면 전월 대비 {abs(change):.1f}% {direction}하였습니다."
            )

    item_data = summary.get("byItem", [])
    if item_data:
        top_item = item_data[0]
        lines.append(
            f"항목별로는 '{top_item.get('항목')}'이(가) "
            f"{top_item.get('총금액', 0):,.0f}원으로 최다 매출 항목입니다."
        )

    return " ".join(lines)
